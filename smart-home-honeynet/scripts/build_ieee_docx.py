#!/usr/bin/env python3
"""Generate an IEEE-conference-formatted Word (.docx) version of the report.

Produces a two-column IEEE paper (Times New Roman, 10pt body, spanning
title/author block, numbered sections and formatted tables) at
docs/report/SmartHoneyNet-IEEE.docx.

Usage:  python3 scripts/build_ieee_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "report" / "SmartHoneyNet-IEEE.docx"


# --------------------------------------------------------------------------
# low-level helpers
# --------------------------------------------------------------------------
def set_cols(section, num: int, space_twips: int = 360) -> None:
    """Set the number of newspaper-style columns on a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def base_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0


def para(doc, text="", *, size=10, bold=False, italic=False, align=None,
         space_after=6, space_before=0, justify=True, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        add_runs(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_runs(p, text, *, size=10, bold=False, italic=False, color=None):
    """Add text, honouring simple **bold** and `code` inline markup."""
    import re

    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        b, i, mono = bold, italic, False
        content = tok
        if tok.startswith("**") and tok.endswith("**"):
            b, content = True, tok[2:-2]
        elif tok.startswith("`") and tok.endswith("`"):
            mono, content = True, tok[1:-1]
        r = p.add_run(content)
        r.bold = b
        r.italic = i
        r.font.size = Pt(size)
        r.font.name = "Consolas" if mono else "Times New Roman"
        if color:
            r.font.color.rgb = color


def heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number}. {title.upper()}")
    r.bold = True
    r.font.size = Pt(10)
    return p


def subheading(doc, letter, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{letter}. {title}")
    r.italic = True
    r.font.size = Pt(10)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text, size=10)
    return p


def table(doc, headers, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(8)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            rr = cells[j].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(8)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(8)
        cap.paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


# --------------------------------------------------------------------------
# document build
# --------------------------------------------------------------------------
def build() -> None:
    doc = Document()
    base_font(doc)

    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(48)
    sec.right_margin = Pt(48)

    # --- single-column title / author block ---
    set_cols(sec, 1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Hunt, Detect and Analyse Attacks on a Smart-Home "
                       "Network using Honeypots, an IDS and the ELK Stack")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Times New Roman"
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("A reproducible platform for smart-home network security · University of Jeddah (CCSE)")
    sr.italic = True
    sr.font.size = Pt(11)
    sub.paragraph_format.space_after = Pt(10)

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ap.add_run("Aleen Saleh Aljohani")
    run.font.size = Pt(11)
    ap.add_run("\naleensaljohani@gmail.com").font.size = Pt(9)
    ap.paragraph_format.space_after = Pt(2)

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = aff.add_run("Cybersecurity Department, College of Computer Science and "
                     "Engineering (CCSE), University of Jeddah, Jeddah, Saudi Arabia")
    ar.font.size = Pt(9)
    aff.paragraph_format.space_after = Pt(10)

    # --- switch to two columns for the body ---
    new = doc.add_section(WD_SECTION.CONTINUOUS)
    new.top_margin = Pt(54)
    new.bottom_margin = Pt(54)
    new.left_margin = Pt(48)
    new.right_margin = Pt(48)
    set_cols(new, 2)

    # Abstract
    ab = doc.add_paragraph()
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abr = ab.add_run("Abstract—")
    abr.bold = True
    abr.italic = True
    abr.font.size = Pt(9)
    abtext = (
        "Smart homes are projected to reach roughly 312 million households by "
        "2027, expanding a poorly-defended IoT attack surface that botnets such "
        "as Mirai exploit through default credentials and exposed telnet/HTTP "
        "interfaces. This paper presents SmartHoneyNet, a reproducible, tested "
        "and quantitatively evaluated platform that combines honeypots, an IDS "
        "and the ELK stack to hunt, detect and analyse these attacks on a "
        "single host. Two low-interaction honeypots emulate a Mirai-style telnet "
        "IoT hub and a smart-home device web panel; a hybrid signature-plus-"
        "anomaly detection engine classifies six attack categories and maps "
        "every alert to MITRE ATT&CK; a pipeline persists all data to both "
        "SQLite and Elasticsearch for Kibana visualisation. The platform runs "
        "on a single host, ships with a 39-test automated suite and enforced "
        "security controls, and on a controlled labelled evaluation set the "
        "engine achieves 100% precision and recall with zero false positives "
        "on benign traffic."
    )
    abr2 = ab.add_run(abtext)
    abr2.italic = True
    abr2.font.size = Pt(9)
    ab.paragraph_format.space_after = Pt(6)

    kw = doc.add_paragraph()
    kwr = kw.add_run("Index Terms—")
    kwr.bold = True
    kwr.italic = True
    kwr.font.size = Pt(9)
    kwr2 = kw.add_run("Honeypot, IoT security, Intrusion Detection, ELK Stack, "
                      "Mirai, MITRE ATT&CK, Smart Home, Threat Hunting.")
    kwr2.italic = True
    kwr2.font.size = Pt(9)
    kw.paragraph_format.space_after = Pt(8)

    # I. INTRODUCTION
    heading(doc, "I", "Introduction")
    para(doc, "Interconnected “smart” devices—thermostats, cameras, "
              "locks, plugs and hubs—have woven themselves into the modern "
              "home. Each device is a small networked computer, and collectively "
              "they form a network whose weakest member sets the security of the "
              "whole. The defensive problem is asymmetric: a homeowner has "
              "neither the tooling nor the expertise of an enterprise SOC, yet "
              "faces the same internet-wide automated adversaries.")
    para(doc, "The word “Hunt” captures the intended posture: not merely "
              "to block attacks reactively, but to attract, observe and "
              "understand them. A honeypot deliberately exposes an attractive, "
              "instrumented target; whatever an attacker does to it is malicious "
              "by definition and worth studying. Pairing honeypots with an IDS "
              "and a log-analytics stack turns those observations into "
              "detections and, ultimately, into intelligence.")
    subheading(doc, "A", "Problem statement")
    para(doc, "Smart-home networks are attacked continuously by automated "
              "botnets and scanners, but homeowners lack visibility into these "
              "attacks and researchers lack reproducible platforms for studying "
              "them in a smart-home context.")
    subheading(doc, "B", "Aim and objectives")
    para(doc, "**Aim.** To build a reproducible platform that hunts, detects and "
              "analyses attacks against a smart-home network, and to evaluate "
              "its detection effectiveness. **Objectives:** (1) emulate "
              "vulnerable devices with safe honeypots; (2) implement a real-time "
              "hybrid IDS; (3) integrate an ELK pipeline and an offline "
              "database; (4) measure detection performance quantitatively; and "
              "(5) enforce responsible-use controls in code.")
    subheading(doc, "C", "Platform capabilities")
    table(doc, ["#", "Capability", "What it delivers"], [
        ["1", "Deployment", "One-command Docker + pure-Python mode"],
        ["2", "Detection coverage", "6 detection categories, each tested"],
        ["3", "Evaluation", "Quantitative precision/recall/F1"],
        ["4", "Detection engine", "Hybrid signature + anomaly"],
        ["5", "Threat taxonomy", "MITRE ATT&CK mapping per alert"],
        ["6", "Quality", "39 tests + CI on 3 Python versions"],
        ["7", "Safe tooling", "Bounded, safety-gated simulators"],
        ["8", "Security", "Threat model + enforced controls"],
    ], caption="TABLE I. Platform capabilities at a glance.")

    # II. BACKGROUND
    heading(doc, "II", "Background")
    subheading(doc, "A", "Smart-home networks and their risks")
    para(doc, "A smart home links appliances and sensors to a hub/router, "
              "controlled through a web or mobile interface. Representative risks "
              "include data and identity theft from unprotected devices, "
              "eavesdropping on device traffic, and (distributed) "
              "denial-of-service amplified by the number of weakly-protected IoT "
              "devices available for recruitment.")
    subheading(doc, "B", "Honeypots")
    para(doc, "A honeypot is an intentionally exposed, instrumented system whose "
              "only purpose is to be attacked. Low-interaction honeypots (used "
              "here) emulate just enough of a service to capture intent while "
              "never executing attacker input—a deliberate safety property. "
              "High-interaction honeypots expose real systems for richer data at "
              "much higher risk.")
    subheading(doc, "C", "Intrusion Detection Systems")
    para(doc, "Signature-based detection matches known-bad patterns—precise "
              "for known threats but blind to novel ones. Anomaly-based "
              "detection models normal behaviour and flags deviations—able "
              "to catch new attacks at the cost of more false positives. "
              "SmartHoneyNet combines both.")
    subheading(doc, "D", "The ELK stack and IoT botnets")
    para(doc, "Elasticsearch stores and indexes events; Logstash ingests and "
              "enriches; Kibana visualises. Mirai and its descendants scan for "
              "open telnet (TCP 23) and try factory-default credentials such as "
              "`root:xc3511` and `admin:admin`—the behaviour that motivates "
              "an accurate telnet honeypot and a default-credential signature.")

    # III. RELATED WORK
    heading(doc, "III", "Related Work")
    para(doc, "Five prior studies frame this work; each leaves a gap this "
              "project addresses.")
    table(doc, ["Ref", "Yr", "Approach", "Limitation"], [
        ["[10]", "2021", "ML attack classification on IoT (Bot-IoT)", "No real-time/streaming eval"],
        ["[11]", "2021", "Threat hunting with Elastic stack", "Needs human input + controls"],
        ["[12]", "2021", "Honeynet + flow classification", "Limited algorithm comparison"],
        ["[13]", "2020", "IoT + Docker (Cowrie) honeypots", "Not all IoT attacks covered"],
        ["[14]", "2019", "Malicious-event detection, ELK+CTI", "False positives; manual cleanup"],
    ], caption="TABLE II. Comparative summary of related studies.")
    para(doc, "**Difference from existing work.** SmartHoneyNet is distinguished "
              "by a hybrid detection engine, an explicit MITRE ATT&CK "
              "classification of every alert, a fully reproducible and "
              "automatically tested implementation, and a quantitative "
              "evaluation with published precision/recall—addressing the "
              "“qualitative only” and “not reproducible” gaps "
              "common to the surveyed work.")

    # IV. THREAT MODEL
    heading(doc, "IV", "Threat Model")
    para(doc, "**Assets:** device credentials, camera/sensor feeds, device "
              "compute and bandwidth, and the LAN foothold. **Adversaries:** "
              "automated IoT botnets (Mirai-class), opportunistic internet-wide "
              "scanners, and manual attackers. **Trust boundary:** the honeynet "
              "lives on an isolated segment that is untrusted inbound and has no "
              "egress; ELK management binds to localhost. Attacker goals are "
              "expressed directly as MITRE ATT&CK techniques, each paired with "
              "the rule that detects it (Table III).")

    # V. METHODOLOGY
    heading(doc, "V", "Methodology")
    para(doc, "The platform follows a three-stage flow, expressed "
              "as software: (1) attack generation—by real attackers or the "
              "safety-gated simulators; (2) detection and collection—"
              "honeypots capture each interaction as a structured event that the "
              "pipeline persists and runs through the IDS; (3) transformation "
              "and analysis—events and alerts are stored in SQLite and "
              "shipped through Logstash into Elasticsearch for Kibana, with an "
              "analyzer producing Markdown/JSON reports.")

    # VI. SYSTEM DESIGN
    heading(doc, "VI", "System Design")
    subheading(doc, "A", "Architecture")
    para(doc, "The platform is a pipeline of composable components that replace "
              "a traditional multi-VM lab with software: the attack simulators "
              "generate adversary traffic, the telnet and HTTP honeypots act as "
              "the sensors, and the detection engine (with an equivalent Suricata "
              "rule file) plus the ELK services perform detection and analysis.")
    subheading(doc, "B", "Detection rules and ATT&CK mapping")
    table(doc, ["Rule", "Type", "MITRE", "Trigger"], [
        ["IOT-DEFAULT-CREDS", "sig", "T1078.001", "Known IoT default credential pair"],
        ["TELNET-BRUTE-FORCE", "anomaly", "T1110.001", "≥5 failed logins / 60 s"],
        ["HTTP-DOS-FLOOD", "anomaly", "T1498", "≥50 requests / 10 s"],
        ["PORT-SCAN", "anomaly", "T1046", "≥8 distinct ports / 30 s"],
        ["SUSPICIOUS-HTTP-PATH", "sig", "T1595.002", "IoT exploit URL (boaform, HNAP1)"],
        ["CMD-INJECTION", "sig", "T1059", "Shell payload (;wget, busybox)"],
    ], caption="TABLE III. Detection rules mapped to MITRE ATT&CK.")
    para(doc, "A per-(source, rule) cooldown collapses a sustained campaign into "
              "a single actionable alert instead of thousands of duplicates.")
    subheading(doc, "C", "Data model and storage")
    para(doc, "Two records flow through the system—Event (an observation) "
              "and Alert (a detection). They serialise identically to a SQLite "
              "row and to the JSON Elasticsearch indexes, keeping the offline "
              "and ELK paths consistent. SQLite is the local system of record; "
              "Elasticsearch provides the same at scale with Kibana dashboards.")

    # VII. IMPLEMENTATION
    heading(doc, "VII", "Implementation")
    para(doc, "The core is implemented in Python using only the standard library "
              "(PyYAML optional), making it portable, auditable and easy to test.")
    bullet(doc, "**Telnet honeypot**—a threaded server presenting a BusyBox "
                "banner and login prompts; records credential pairs, “accepts"
                "” weak credentials to lure the attacker into a fake shell, "
                "and logs typed commands. Nothing is ever executed.")
    bullet(doc, "**HTTP honeypot**—emulates a “SmartHome Hub 2.4” "
                "admin panel; records method, path, user-agent and POSTed "
                "credentials, caps body size and returns static HTML.")
    bullet(doc, "**Detection engine**—per-source sliding-window deques for "
                "rate rules and compiled regex sets for signatures; thread-safe "
                "and stream-oriented.")
    bullet(doc, "**Attack simulators**—port scan, brute-force and a bounded, "
                "safety-gated HTTP flood, each checked by the safety module "
                "before any packet is sent.")
    para(doc, "`docker compose up -d --build` starts the honeypots plus "
              "Elasticsearch, Logstash and Kibana on an isolated bridge network; "
              "a saved-objects file provisions the dashboards. A Suricata rule "
              "set mirrors the engine for network-layer detection.")

    # VIII. EXPERIMENTS AND EVALUATION
    heading(doc, "VIII", "Experiments and Evaluation")
    subheading(doc, "A", "End-to-end demonstration")
    para(doc, "The demo command starts both honeypots on ephemeral ports and "
              "launches a port scan, a telnet brute-force, an HTTP flood and "
              "web-exploitation probes against them. A representative run "
              "captured ~196 events and raised alerts in all five exercised "
              "categories—covering six attack types end to end, as a single "
              "reproducible command.")
    subheading(doc, "B", "Quantitative detection evaluation")
    para(doc, "A labelled event stream (94 events) interleaves benign, "
              "human-paced traffic with four attack campaigns from distinct "
              "sources with known ground-truth categories. Metrics are computed "
              "at source-campaign granularity, matching how the rate rules fire.")
    table(doc, ["Category", "P", "R", "F1"], [
        ["default-credentials", "1.00", "1.00", "1.00"],
        ["brute-force", "1.00", "1.00", "1.00"],
        ["denial-of-service", "1.00", "1.00", "1.00"],
        ["active-scanning", "1.00", "1.00", "1.00"],
        ["command-injection", "1.00", "1.00", "1.00"],
        ["Micro-average", "1.00", "1.00", "1.00"],
    ], caption="TABLE IV. Detection performance on the labelled set "
               "(2 benign sources, 0 false positives).")
    para(doc, "The engine detected every attack campaign in its correct category "
              "and raised no alerts on benign traffic in this controlled set; "
              "the result is asserted by a regression test. Section X is candid "
              "about why real-world traffic is harder.")
    subheading(doc, "C", "Automated test suite")
    para(doc, "39 pytest tests cover the database, every detection rule with "
              "synthetic timelines, live honeypot interaction, the safety "
              "guardrails, the pipeline, the analyzer, the CLI and a full "
              "integration run; line coverage is ~84%. CI runs the suite plus "
              "the demo on Python 3.9, 3.11 and 3.12.")

    # IX. RESULTS AND DISCUSSION
    heading(doc, "IX", "Results and Discussion")
    para(doc, "The experiments confirm the central hypothesis and strengthen it "
              "with measurement. Three findings stand out: (1) default "
              "credentials remain the highest-signal indicator—a single "
              "login with a Mirai default pair is immediately malicious; (2) "
              "hybrid detection is complementary, not redundant—signatures "
              "catch the known while rate rules catch behaviour that has no "
              "fixed signature; and (3) the captured credential intelligence has "
              "direct defensive value, yielding a blocklist of passwords a "
              "homeowner must never use.")

    # X. SECURITY, ETHICS AND LIMITATIONS
    heading(doc, "X", "Security, Ethics and Limitations")
    para(doc, "Because the project ships attack tooling, ethics are enforced in "
              "code: simulators call a safety check that refuses any non-private "
              "target; the DoS simulator enforces hard caps; honeypots are "
              "low-interaction with bounded buffers and timeouts; containers run "
              "non-root on an isolated network.")
    para(doc, "**Limitations.** The honeypots emulate telnet/HTTP interfaces, "
              "not full firmware or radios (Zigbee/Z-Wave/BLE). The perfect "
              "scores are on a clean labelled set; production traffic would "
              "produce false positives and require threshold tuning. Rate rules "
              "need a warm-up window and can miss low-and-slow attacks. These "
              "are the honest boundaries within which the results hold.")

    # XI. FUTURE WORK
    heading(doc, "XI", "Future Work")
    para(doc, "Future directions include Cowrie/Dionaea integration for "
              "higher-interaction capture; an MQTT honeypot (TCP 1883) for the "
              "dominant smart-home broker; a virtual IoT-device environment with "
              "realistic device fingerprints; an ML-based anomaly detector "
              "evaluated on the same labelled sets; GeoIP and threat-intel "
              "enrichment; and automated firewall response for high-severity "
              "sources.")

    # XII. CONCLUSION
    heading(doc, "XII", "Conclusion")
    para(doc, "SmartHoneyNet is a reproducible, tested and measured platform for "
              "smart-home network security. It builds on a proven core idea—"
              "honeypots to attract, an IDS to detect, ELK to analyse—and adds a hybrid engine "
              "covering six attack classes, a MITRE ATT&CK classification of "
              "every alert, a real database and dashboards, a 39-test suite with "
              "CI, enforced security controls, and a quantitative evaluation "
              "showing 100% precision and recall with zero false positives on a "
              "controlled labelled set. As smart-home adoption races toward "
              "hundreds of millions of households, tools that let defenders hunt "
              "attackers safely, reproducibly and measurably are exactly what "
              "the field needs.")

    # References
    heading(doc, "", "References")
    refs = [
        "J. Lasquety-Reyes, “Global: Smart home—number of users 2018–2027,” Statista.",
        "“What is a smart home and what are the benefits?,” Constellation.com.",
        "“Smart home: Threats and countermeasures,” Rambus, 2017.",
        "“A Study of Smart Home Environment and its Security Threats,” ResearchGate.",
        "“Smart home: Threats and countermeasures,” Rambus, 2017.",
        "K. Subramanian and W. Meng, “Threat Hunting Using Elastic Stack: An Evaluation,” IEEE SOLI, 2021.",
        "H. Qu, D. Qu and J. Tong, “Virtual environment for smart home,” IEEE CYBER, 2015.",
        "“What is an intrusion detection system?,” DNSstuff, 2019.",
        "E. Hasson and L. Cheng, “Honeypot,” Imperva Learning Center.",
        "A. Churcher et al., “An experimental analysis of attack classification using ML in IoT networks,” Sensors, 21(2):446, 2021.",
        "K. Subramanian and W. Meng, “Threat Hunting Using Elastic Stack,” IEEE SOLI, 2021.",
        "Banerjee et al., “Network traffic analysis based IoT botnet detection using honeynet data,” 2021.",
        "S. Bistarelli, E. Bosimini and F. Santini, “A Report on the Security of Home Connections with IoT and Docker Honeypots,” CEUR-WS Vol-2597.",
        "M. Harikanth and P. Rajarajeswari, “Malicious event detection using ELK stack through CTI,” IJITEE, 8(7):882–886, 2019.",
        "MITRE ATT&CK® Enterprise Matrix, The MITRE Corporation.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"[{i}] {r}")
        run.font.size = Pt(8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
